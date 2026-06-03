import os
import time
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager

print("Word Robotu hazırlanıyor devrem, kemerlerini bağla...")

ana_klasor = os.path.dirname(os.path.abspath(__file__))

# Yeni bir Word belgesi oluşturuyoruz
doc = Document()
doc.add_heading('Yargıtay Kararları', 0)

ayarlar = Options()
servis = Service(ChromeDriverManager().install())
tarayici = webdriver.Chrome(service=servis, options=ayarlar)
tarayici.get("https://karararama.yargitay.gov.tr/")

print("\n" + "="*60)
print("SİSTEM HAZIR VE BEKLEMEDE!")
print("1. Chrome'da aramanı yap.")
print("2. Kararlar ekranda listelendiğinde bu siyah ekrana gelip ENTER'a bas!")
print("="*60 + "\n")

input("Kopyalama işlemini başlatmak için ENTER'a bas devrem...")
print("Operasyon başlıyor! Hedef: Metinleri saniyeler içinde Word'e çekmek.")

try:
    try:
        dropdown_element = tarayici.find_element(By.XPATH, "//select[contains(@name, 'length') or contains(@class, 'custom-select')]")
        Select(dropdown_element).select_by_visible_text("100")
        print("Sayfa başı 100 karar gösterimi seçildi...")
        time.sleep(3)
    except:
        print("Mevcut liste üzerinden devam ediliyor.")

    satirlar = tarayici.find_elements(By.XPATH, "//table//tbody/tr")
    toplam_karar = len(satirlar)
    
    if toplam_karar == 0:
        print("Uyarı: Listede karar bulunamadı!")
    else:
        print(f"Listede {toplam_karar} adet karar tespit edildi. Aktarım başlıyor...\n")
        
        basarili_islem = 0
        for i in range(toplam_karar):
            try:
                guncel_satirlar = tarayici.find_elements(By.XPATH, "//table//tbody/tr")
                secili_satir = guncel_satirlar[i]
                
                # Satıra tıkla
                tarayici.execute_script("arguments[0].scrollIntoView({block: 'center'});", secili_satir)
                time.sleep(0.3)
                secili_satir.click()
                
                # Kararın sağda açılması için bekle
                time.sleep(2)
                
                # Sayfanın tüm metnini al
                sayfa_metni = tarayici.find_element(By.TAG_NAME, "body").text
                
                # Sadece karar metnini filtrele (İçtihat Metni ile başlar, Önceki/Sonraki Karar butonlarıyla biter)
                bolunmus = sayfa_metni.split("İçtihat Metni")
                
                if len(bolunmus) > 1:
                    # Metni ayıkla
                    saf_karar = "İçtihat Metni" + bolunmus[-1]
                    
                    # Alttaki gereksiz buton yazılarını temizle
                    if "Önceki Karar" in saf_karar:
                        saf_karar = saf_karar.split("Önceki Karar")[0]
                    elif "Sonraki Karar" in saf_karar:
                        saf_karar = saf_karar.split("Sonraki Karar")[0]
                        
                    # Word belgesine başlık ve metin olarak ekle
                    doc.add_heading(f"{i+1}. Karar", level=1)
                    doc.add_paragraph(saf_karar.strip())
                    doc.add_page_break() # Her karar yeni sayfada başlasın
                    
                    basarili_islem += 1
                    print(f"[{basarili_islem}/{toplam_karar}] Karar Word'e başarıyla yazıldı.")
                else:
                    print(f"[{i+1}. Karar] atlandı: Ekranda metin bulunamadı.")
                    
            except Exception as e:
                print(f"[{i+1}. Karar] atlandı.")

except Exception as e:
    print(f"Sitede beklenmeyen hata: {e}")

tarayici.quit()

# Word dosyasını kaydet
sonuc_yolu = os.path.join(ana_klasor, "Toplu_Yargitay_Kararlari.docx")
doc.save(sonuc_yolu)

print(f"\nZAFER! {basarili_islem} adet karar tek bir Word dosyasına kusursuzca aktarıldı devrem.")
print(f"Dosyanı şurada bulabilirsin: {sonuc_yolu}")